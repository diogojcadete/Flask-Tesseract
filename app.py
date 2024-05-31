import os
from flask import Flask, request, render_template, send_from_directory
from werkzeug.utils import secure_filename
import pytesseract
from PIL import Image
from tesserocr import PyTessBaseAPI, RIL, iterate_level
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['RESULT_FOLDER'] = 'results'
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg'}

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

if not os.path.exists(app.config['RESULT_FOLDER']):
    os.makedirs(app.config['RESULT_FOLDER'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def get_font(image_path):
    with PyTessBaseAPI(path='tessdata-main') as api:
        api.SetImageFile(image_path)
        api.Recognize()
        ri = api.GetIterator()
        level = RIL.SYMBOL
        counter = 0
        total_font_size = 0

        for r in iterate_level(ri, level):
            symbol = r.GetUTF8Text(level)
            word_attributes = r.WordFontAttributes()

            if symbol:
                total_font_size += word_attributes['pointsize']
                counter += 1
                print(u'symbol {}, font: {}, size {}'.format(symbol, word_attributes['font_name'], word_attributes['pointsize']))

        if counter > 0:
            average_font_size = total_font_size / counter
            return average_font_size
        else:
            return None

def ocr_to_files(image_path, output_html_path, output_txt_path, output_docx_path):
    extracted_text = pytesseract.image_to_string(Image.open(image_path))

    average_font_size = get_font(image_path)

    html_content = f"<html><head><title></title><style>pre {{ font-size: {average_font_size}px; }}</style></head><body><pre>{extracted_text}</pre></body></html>"

    with open(output_html_path, "w") as html_file:
        html_file.write(html_content)

    with open(output_txt_path, "w") as txt_file:
        txt_file.write(extracted_text)

    doc = Document()
    for paragraph in extracted_text.split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(paragraph)
        font = run.font
        font.size = Pt(average_font_size)
    doc.save(output_docx_path)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return 'No file part'
        file = request.files['file']
        if file.filename == '':
            return 'No selected file'
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(image_path)
            
            base_filename = os.path.splitext(filename)[0]
            output_html_path = os.path.join(app.config['RESULT_FOLDER'], base_filename + '.html')
            output_txt_path = os.path.join(app.config['RESULT_FOLDER'], base_filename + '.txt')
            output_docx_path = os.path.join(app.config['RESULT_FOLDER'], base_filename + '.docx')

            ocr_to_files(image_path, output_html_path, output_txt_path, output_docx_path)

            return render_template('results.html', filename=base_filename)
    return render_template('upload.html')

@app.route('/results/<filename>')
def download_file(filename):
    return send_from_directory(app.config['RESULT_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)
